"""
Generate Word Report with Screenshots for Water Potability ML Project
Author: Sohaib Farooq
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_heading('Water Potability Prediction System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Using Machine Learning for Safe Drinking Water Classification')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author info
author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_para.add_run('Sohaib Farooq\nsohaib.farooq@bigacademy.com\nMSc Computing - Independent Project')
author_run.font.size = Pt(12)
author_run.font.italic = True

doc.add_paragraph()

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'This study presents a comprehensive machine learning approach for predicting water potability '
    'based on physicochemical parameters. Access to safe drinking water is a fundamental human right, '
    'and accurate prediction of water safety is critical for public health. We implemented and compared '
    'four machine learning algorithms: Logistic Regression, Random Forest, XGBoost, and LightGBM. '
    'The dataset comprises 3,276 water samples with nine quality indicators. Our analysis reveals that '
    'Random Forest achieves the highest accuracy of 66.01%, while LightGBM demonstrates the best '
    'balance between precision and recall. A web-based application was developed using Streamlit to '
    'enable real-time water potability prediction.'
)

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'Water quality assessment is essential for ensuring public health and environmental sustainability. '
    'Traditional laboratory-based water testing methods, while accurate, are time-consuming and expensive. '
    'Machine learning offers a promising alternative for rapid, cost-effective water quality prediction.'
)
doc.add_paragraph(
    'This project addresses the binary classification problem of determining whether water is potable '
    '(safe for human consumption) or non-potable based on nine physicochemical parameters. The objective '
    'is to develop and compare multiple machine learning models to identify the most effective approach '
    'for water potability prediction.'
)

# Methodology
doc.add_heading('2. Methodology', level=1)

doc.add_heading('2.1 Dataset Description', level=2)
doc.add_paragraph(
    'The Water Potability dataset from Kaggle contains 3,276 water samples with the following features:'
)

# Features table
table = doc.add_table(rows=10, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Unit'

features = [
    ('pH', 'Acidity/Alkalinity level', '0-14'),
    ('Hardness', 'Capacity to precipitate soap', 'mg/L'),
    ('Solids', 'Total dissolved solids', 'ppm'),
    ('Chloramines', 'Amount of chloramines', 'ppm'),
    ('Sulfate', 'Sulfate dissolved in water', 'mg/L'),
    ('Conductivity', 'Electrical conductivity', 'uS/cm'),
    ('Organic_carbon', 'Organic carbon amount', 'ppm'),
    ('Trihalomethanes', 'THMs amount', 'ug/L'),
    ('Turbidity', 'Cloudiness measure', 'NTU'),
]

for i, (feat, desc, unit) in enumerate(features, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = feat
    row_cells[1].text = desc
    row_cells[2].text = unit

doc.add_paragraph()

doc.add_heading('2.2 Machine Learning Algorithms', level=2)

# Logistic Regression
doc.add_heading('2.2.1 Logistic Regression', level=3)
doc.add_paragraph(
    'Logistic Regression is a linear model for binary classification. The probability of class membership '
    'is modeled using the sigmoid function:'
)
doc.add_paragraph('P(y=1|x) = 1 / (1 + exp(-(w^T x + b)))')
doc.add_paragraph(
    'To address class imbalance in the dataset, we applied class_weight="balanced" which adjusts '
    'weights inversely proportional to class frequencies.'
)

# Random Forest
doc.add_heading('2.2.2 Random Forest', level=3)
doc.add_paragraph(
    'Random Forest is an ensemble method that constructs multiple decision trees during training. '
    'The final prediction is determined by majority voting:'
)
doc.add_paragraph('y_hat = mode{h_1(x), h_2(x), ..., h_B(x)}')
doc.add_paragraph('where h_b represents individual decision trees and B is the number of trees (100 in our implementation).')

# XGBoost
doc.add_heading('2.2.3 XGBoost', level=3)
doc.add_paragraph(
    'XGBoost (Extreme Gradient Boosting) is a gradient boosting algorithm that optimizes:'
)
doc.add_paragraph('L(phi) = sum[l(y_i, y_hat_i)] + sum[Omega(f_k)]')
doc.add_paragraph(
    'where l is the loss function and Omega is the regularization term. We used 200 estimators with '
    'learning rate of 0.1 and max depth of 5.'
)

# LightGBM
doc.add_heading('2.2.4 LightGBM', level=3)
doc.add_paragraph(
    'LightGBM uses histogram-based algorithms for efficient tree construction. It employs '
    'Gradient-based One-Side Sampling (GOSS) and Exclusive Feature Bundling (EFB) for faster training.'
)

doc.add_heading('2.3 Evaluation Metrics', level=2)
doc.add_paragraph('The models were evaluated using the following metrics:')
doc.add_paragraph('- Accuracy = (TP + TN) / (TP + TN + FP + FN)')
doc.add_paragraph('- Precision = TP / (TP + FP)')
doc.add_paragraph('- Recall = TP / (TP + FN)')
doc.add_paragraph('- F1-Score = 2 * (Precision * Recall) / (Precision + Recall)')
doc.add_paragraph('- AUC-ROC: Area under the Receiver Operating Characteristic curve')

# Web Application Section
doc.add_heading('3. Web Application', level=1)
doc.add_paragraph(
    'A Streamlit-based web application was developed to provide an interactive interface for water '
    'potability prediction. The application includes the following features:'
)
doc.add_paragraph('- Home Page: Project overview and quick prediction capability')
doc.add_paragraph('- Data Exploration: Dataset visualization with distribution charts and correlation heatmap')
doc.add_paragraph('- Model Training: Display of training results and confusion matrices')
doc.add_paragraph('- Prediction: Interactive sliders for inputting water quality parameters')
doc.add_paragraph('- Model Comparison: Side-by-side comparison of all algorithms')

# Add screenshots
screenshot_dir = '.playwright-mcp'

doc.add_heading('3.1 Application Screenshots', level=2)

screenshots = [
    ('01_home_page.png', 'Figure 1: Home Page - Water Potability Prediction System'),
    ('02_data_distributions.png', 'Figure 2: Data Exploration - Target Variable Distribution'),
    ('03_correlation_heatmap.png', 'Figure 3: Feature Correlation Heatmap'),
    ('04_model_training.png', 'Figure 4: Model Training Results'),
    ('05_model_comparison.png', 'Figure 5: Model Comparison with Accuracy Chart'),
    ('06_prediction_page.png', 'Figure 6: Prediction Interface'),
    ('07_prediction_result.png', 'Figure 7: Prediction Result with Confidence Scores'),
]

for filename, caption in screenshots:
    filepath = os.path.join(screenshot_dir, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(5.5))
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].font.italic = True
        caption_para.runs[0].font.size = Pt(10)
        doc.add_paragraph()

# Results
doc.add_heading('4. Results', level=1)

doc.add_heading('4.1 Model Performance Comparison', level=2)

# Results table
results_table = doc.add_table(rows=5, cols=5)
results_table.style = 'Table Grid'

headers = ['Algorithm', 'Accuracy', 'Precision', 'Recall', 'F1-Score']
for i, header in enumerate(headers):
    results_table.rows[0].cells[i].text = header

results_data = [
    ('Logistic Regression', '52.59%', '0.4159', '0.5312', '0.4666'),
    ('Random Forest', '66.01%', '0.6514', '0.2773', '0.3890'),
    ('XGBoost', '61.89%', '0.5134', '0.4492', '0.4792'),
    ('LightGBM', '62.20%', '0.5174', '0.4648', '0.4897'),
]

for i, row_data in enumerate(results_data, 1):
    for j, cell_data in enumerate(row_data):
        results_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

doc.add_heading('4.2 Key Findings', level=2)
doc.add_paragraph(
    '1. Random Forest achieved the highest accuracy (66.01%) and precision (0.6514), making it '
    'the best model for minimizing false positives.'
)
doc.add_paragraph(
    '2. Logistic Regression with balanced class weights achieved the highest recall (0.5312), '
    'indicating better detection of potable water samples.'
)
doc.add_paragraph(
    '3. LightGBM demonstrated the best F1-Score (0.4897), showing the most balanced trade-off '
    'between precision and recall.'
)
doc.add_paragraph(
    '4. The class imbalance in the dataset (61% non-potable vs 39% potable) significantly impacts '
    'model performance, particularly affecting recall scores.'
)

# Discussion
doc.add_heading('5. Discussion', level=1)
doc.add_paragraph(
    'The results indicate that ensemble methods (Random Forest, XGBoost, LightGBM) generally '
    'outperform Logistic Regression in terms of accuracy. However, the choice of the best model '
    'depends on the specific use case:'
)
doc.add_paragraph(
    '- For applications where false positives are costly (incorrectly labeling unsafe water as safe), '
    'Random Forest is recommended due to its high precision.'
)
doc.add_paragraph(
    '- For applications where missing potable water is costly (false negatives), '
    'Logistic Regression with balanced weights provides better recall.'
)
doc.add_paragraph(
    '- For a balanced approach, LightGBM offers the best F1-Score with competitive accuracy.'
)

# Conclusion
doc.add_heading('6. Conclusion', level=1)
doc.add_paragraph(
    'This study successfully implemented and compared four machine learning algorithms for water '
    'potability prediction. The Random Forest classifier emerged as the best-performing model with '
    '66.01% accuracy. A user-friendly web application was developed to enable real-time predictions, '
    'making the technology accessible to non-technical users.'
)
doc.add_paragraph(
    'Future work could explore deep learning approaches, feature engineering techniques, and '
    'ensemble methods combining multiple classifiers to further improve prediction accuracy.'
)

# References
doc.add_heading('7. References', level=1)
doc.add_paragraph('[1] Kaggle. Water Potability Dataset. https://www.kaggle.com/datasets/adityakadiwal/water-potability')
doc.add_paragraph('[2] Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. JMLR 12, pp. 2825-2830.')
doc.add_paragraph('[3] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. KDD 2016.')
doc.add_paragraph('[4] Ke, G., et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision Tree. NIPS 2017.')
doc.add_paragraph('[5] World Health Organization. Guidelines for Drinking-water Quality. https://www.who.int/')

# Appendix
doc.add_heading('Appendix A: Deployed Application', level=1)
doc.add_paragraph('The web application is deployed and accessible at:')
doc.add_paragraph('https://water-potability-prediction-ip-assignment.streamlit.app/')
doc.add_paragraph()
doc.add_paragraph('Source Code Repository:')
doc.add_paragraph('https://github.com/sohaib3335/water-potability-prediction')

# Save document
doc.save('Water_Potability_ML_Report.docx')
print('Report generated successfully: Water_Potability_ML_Report.docx')
